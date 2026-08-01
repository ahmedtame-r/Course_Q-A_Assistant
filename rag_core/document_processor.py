"""
Document discovery, validation, cleaning and chunking.

``DocumentProcessor`` walks the ``data/<course>/`` folders, loads every
supported file (PDF, DOCX, TXT, CSV), normalises the extracted text and
splits it into overlapping chunks carrying rich metadata that is later
used for source attribution.
"""

from __future__ import annotations

import csv
import json
import re
import unicodedata
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Iterator

from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from rag_core.config import AppConfig
from rag_core.logging_setup import get_logger

try:
    import docx  # python-docx
    if not hasattr(docx, "Document"):
        # The 'docx' and 'python-docx' PyPI packages both install a module
        # literally named 'docx' but expose different APIs. If the wrong
        # one is installed, importing succeeds but Document() is missing.
        docx = None
        _DOCX_IMPORT_ERROR = (
            "A package named 'docx' is installed, but it does not provide "
            "Document(). This usually means the wrong package is installed - "
            "run 'pip uninstall docx' then 'pip install python-docx'."
        )
    else:
        _DOCX_IMPORT_ERROR = None
except ImportError as exc:  # pragma: no cover
    docx = None
    _DOCX_IMPORT_ERROR = f"{exc}. Run 'pip install python-docx' in this environment."


@dataclass
class RawDocument:
    """A single loaded source document before chunking."""

    text: str
    course: str
    file_name: str
    file_path: str
    doc_type: str
    page_number: int | None = None


@dataclass
class ProcessedChunk:
    """A cleaned, size-bounded chunk ready for embedding."""

    chunk_id: str
    text: str
    course: str
    file_name: str
    doc_type: str
    page_number: int | None = None
    char_count: int = field(default=0)

    def __post_init__(self) -> None:
        self.char_count = len(self.text)

    def to_metadata(self) -> dict:
        """Metadata dict suitable for storage alongside the embedding."""
        return {
            "chunk_id": self.chunk_id,
            "course": self.course,
            "file_name": self.file_name,
            "doc_type": self.doc_type,
            "page_number": self.page_number if self.page_number is not None else -1,
            "char_count": self.char_count,
        }


class DocumentProcessor:
    """Loads course materials from disk and turns them into chunks."""

    def __init__(self, config: AppConfig) -> None:
        self.config = config
        self.logger = get_logger(self.__class__.__name__, config.log_level)
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=list(config.chunk_separators),
            length_function=len,
        )

    # ------------------------------------------------------------------
    # Discovery
    # ------------------------------------------------------------------
    def discover_files(self) -> list[Path]:
        """Return every file under ``data/`` with a supported extension."""
        discovered: list[Path] = []
        if not self.config.data_dir.exists():
            self.logger.warning("Data directory does not exist: %s", self.config.data_dir)
            return discovered

        for path in sorted(self.config.data_dir.rglob("*")):
            if not path.is_file():
                continue
            if path.suffix.lower() in self.config.supported_extensions:
                discovered.append(path)
            else:
                self.logger.warning("Skipping unsupported file: %s", path)

        self.logger.info("Discovered %d supported file(s).", len(discovered))
        return discovered

    # ------------------------------------------------------------------
    # Loading
    # ------------------------------------------------------------------
    def load_file(self, path: Path) -> list[RawDocument]:
        """Dispatch to the right loader based on file extension."""
        course = path.parent.name
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf":
                return self._load_pdf(path, course)
            if suffix == ".docx":
                return self._load_docx(path, course)
            if suffix == ".txt":
                return self._load_txt(path, course)
            if suffix == ".csv":
                return self._load_csv(path, course)
        except (PdfReadError, OSError, ValueError) as exc:
            self.logger.error("Failed to read '%s': %s", path, exc)
            return []

        self.logger.warning("No loader registered for extension '%s'", suffix)
        return []

    def _load_pdf(self, path: Path, course: str) -> list[RawDocument]:
        reader = PdfReader(str(path))
        documents: list[RawDocument] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(
                    RawDocument(
                        text=text,
                        course=course,
                        file_name=path.name,
                        file_path=str(path),
                        doc_type="pdf",
                        page_number=page_number,
                    )
                )
        if not documents:
            self.logger.warning("No extractable text found in PDF: %s", path)
        return documents

    def _load_docx(self, path: Path, course: str) -> list[RawDocument]:
        if docx is None:
            raise RuntimeError(_DOCX_IMPORT_ERROR or "python-docx is not installed")

        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        table_lines: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))

        full_text = "\n".join(paragraphs + table_lines)
        if not full_text.strip():
            self.logger.warning("DOCX file has no extractable text: %s", path)
            return []

        return [
            RawDocument(
                text=full_text,
                course=course,
                file_name=path.name,
                file_path=str(path),
                doc_type="docx",
                page_number=None,
            )
        ]

    def _load_txt(self, path: Path, course: str) -> list[RawDocument]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        if not text.strip():
            self.logger.warning("TXT file is empty: %s", path)
            return []

        return [
            RawDocument(
                text=text,
                course=course,
                file_name=path.name,
                file_path=str(path),
                doc_type="txt",
                page_number=None,
            )
        ]

    def _load_csv(self, path: Path, course: str) -> list[RawDocument]:
        with open(path, newline="", encoding="utf-8", errors="ignore") as handle:
            reader = csv.reader(handle)
            rows = list(reader)

        if not rows:
            self.logger.warning("CSV file is empty: %s", path)
            return []

        header, *body = rows
        lines = [", ".join(header)]
        for row in body:
            paired = [f"{col}: {val}" for col, val in zip(header, row)]
            lines.append("; ".join(paired))

        return [
            RawDocument(
                text="\n".join(lines),
                course=course,
                file_name=path.name,
                file_path=str(path),
                doc_type="csv",
                page_number=None,
            )
        ]

    # ------------------------------------------------------------------
    # Cleaning
    # ------------------------------------------------------------------
    @staticmethod
    def clean_text(text: str) -> str:
        """Normalise unicode, collapse whitespace, drop duplicate blank lines."""
        text = unicodedata.normalize("NFKC", text)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        return text.strip()

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------
    def chunk_documents(self, raw_documents: list[RawDocument]) -> list[ProcessedChunk]:
        """Clean and split raw documents into `ProcessedChunk` objects."""
        chunks: list[ProcessedChunk] = []
        chunk_counter = 0

        for raw_document in raw_documents:
            cleaned = self.clean_text(raw_document.text)
            if not cleaned:
                continue

            for piece in self.splitter.split_text(cleaned):
                if not piece.strip():
                    continue
                chunk_counter += 1
                chunk_id = self._build_chunk_id(raw_document, chunk_counter)
                chunks.append(
                    ProcessedChunk(
                        chunk_id=chunk_id,
                        text=piece,
                        course=raw_document.course,
                        file_name=raw_document.file_name,
                        doc_type=raw_document.doc_type,
                        page_number=raw_document.page_number,
                    )
                )

        self.logger.info("Created %d chunk(s) from %d document page(s)/section(s).", len(chunks), len(raw_documents))
        return chunks

    @staticmethod
    def _build_chunk_id(raw_document: RawDocument, counter: int) -> str:
        stem = Path(raw_document.file_name).stem.replace(" ", "_")
        page_part = f"p{raw_document.page_number}" if raw_document.page_number else "p0"
        return f"{raw_document.course.replace(' ', '_')}__{stem}__{page_part}__c{counter:05d}"

    # ------------------------------------------------------------------
    # Pipeline entry point
    # ------------------------------------------------------------------
    def process_all(self) -> list[ProcessedChunk]:
        """Discover, load, clean and chunk every file under ``data/``."""
        all_raw_documents: list[RawDocument] = []
        for path in self.discover_files():
            all_raw_documents.extend(self.load_file(path))

        if not all_raw_documents:
            self.logger.warning("No documents were loaded. Check the data/ directory contents.")

        return self.chunk_documents(all_raw_documents)

    # ------------------------------------------------------------------
    # Persistence helpers
    # ------------------------------------------------------------------
    def save_chunks(self, chunks: list[ProcessedChunk], path: Path | None = None) -> Path:
        """Persist processed chunks to a JSON Lines file."""
        output_path = path or self.config.processed_chunks_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, "w", encoding="utf-8") as handle:
            for chunk in chunks:
                handle.write(json.dumps(asdict(chunk), ensure_ascii=False) + "\n")

        self.logger.info("Saved %d chunk(s) to %s", len(chunks), output_path)
        return output_path

    def load_chunks(self, path: Path | None = None) -> list[ProcessedChunk]:
        """Load previously persisted chunks from a JSON Lines file."""
        input_path = path or self.config.processed_chunks_path
        if not input_path.exists():
            raise FileNotFoundError(f"Processed chunks file not found: {input_path}")

        chunks: list[ProcessedChunk] = []
        with open(input_path, encoding="utf-8") as handle:
            for line in handle:
                if not line.strip():
                    continue
                record = json.loads(line)
                record.pop("char_count", None)
                chunks.append(ProcessedChunk(**record))

        self.logger.info("Loaded %d chunk(s) from %s", len(chunks), input_path)
        return chunks


def iter_course_files(config: AppConfig) -> Iterator[tuple[str, Path]]:
    """Convenience generator yielding ``(course_name, file_path)`` pairs."""
    for course in config.course_names:
        course_dir = config.data_dir / course
        if not course_dir.exists():
            continue
        for path in sorted(course_dir.iterdir()):
            if path.is_file() and path.suffix.lower() in config.supported_extensions:
                yield course, path
